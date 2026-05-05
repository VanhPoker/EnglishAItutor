#!/usr/bin/env python3
"""Build the academic DATN report for EnglishAItutor.

The report intentionally follows the structure of the provided HaUI sample:
front matter, theory chapter, analysis/design chapter, implementation/results
chapter, conclusion, and references. Demo strategy and defense notes are kept
out of this document.
"""

from __future__ import annotations

import math
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "NguyenVietAnh2021600938-DATN-academic.docx"
ASSET_DIR = ROOT / "report_assets_academic"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
UI_DIR = ROOT / "report_assets_final" / "ui_cropped"
BASE_ASSET_DIR = ROOT / "report_assets"

FONT_DIR = Path("/System/Library/Fonts/Supplemental")
FONT = str(FONT_DIR / "Arial.ttf")
FONT_BOLD = str(FONT_DIR / "Arial Bold.ttf")
FONT_ITALIC = str(FONT_DIR / "Arial Italic.ttf")


def font(size: int, bold: bool = False, italic: bool = False):
    path = FONT_BOLD if bold else FONT_ITALIC if italic else FONT
    return ImageFont.truetype(path, size=size)


def wrap_lines(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt) -> list[str]:
    out: list[str] = []
    for para in text.split("\n"):
        if not para:
            out.append("")
            continue
        words = para.split()
        line = ""
        for word in words:
            candidate = word if not line else f"{line} {word}"
            if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
                line = candidate
            else:
                if line:
                    out.append(line)
                line = word
        if line:
            out.append(line)
    return out


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    max_width: int,
    fnt,
    fill="#111827",
    spacing=5,
    align="left",
) -> int:
    x, y = xy
    lines = wrap_lines(draw, text, max_width, fnt)
    line_height = fnt.size + spacing
    for line in lines:
        tx = x
        if align == "center":
            w = draw.textbbox((0, 0), line, font=fnt)[2]
            tx = x + (max_width - w) / 2
        draw.text((tx, y), line, font=fnt, fill=fill)
        y += line_height
    return y


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    subtitle: str = "",
    fill="#F8FAFC",
    outline="#334155",
    title_fill="#111827",
    radius=18,
):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    tf = font(28, bold=True)
    sf = font(20)
    tw = draw.textbbox((0, 0), title, font=tf)[2]
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 24), title, font=tf, fill=title_fill)
    if subtitle:
        draw_wrapped(draw, (x1 + 20, y1 + 70), subtitle, x2 - x1 - 40, sf, fill="#334155", align="center")


def arrow(draw: ImageDraw.ImageDraw, p1: tuple[int, int], p2: tuple[int, int], fill="#2563EB", width=4):
    draw.line([p1, p2], fill=fill, width=width)
    ang = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
    head = 16
    left = (p2[0] - head * math.cos(ang - math.pi / 6), p2[1] - head * math.sin(ang - math.pi / 6))
    right = (p2[0] - head * math.cos(ang + math.pi / 6), p2[1] - head * math.sin(ang + math.pi / 6))
    draw.polygon([p2, left, right], fill=fill)


def make_architecture(path: Path):
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 45), "Kiến trúc tổng thể hệ thống EnglishAItutor", font=font(42, bold=True), fill="#111827")
    boxes = {
        "learner": (90, 250, 380, 390),
        "frontend": (540, 170, 890, 330),
        "api": (540, 470, 890, 630),
        "livekit": (1030, 170, 1380, 330),
        "agent": (1030, 470, 1380, 630),
        "llm": (1450, 470, 1740, 630),
        "db": (540, 770, 890, 930),
        "memory": (1030, 770, 1380, 930),
    }
    rounded_box(d, boxes["learner"], "Học viên", "Trình duyệt web")
    rounded_box(d, boxes["frontend"], "React + Vite", "Giao diện học viên và quản trị", fill="#E0F2FE")
    rounded_box(d, boxes["api"], "FastAPI", "Auth, quiz, billing, token", fill="#FEF3C7")
    rounded_box(d, boxes["livekit"], "LiveKit", "WebRTC room, data channel", fill="#F3E8FF")
    rounded_box(d, boxes["agent"], "Agent Worker", "STT, LangGraph, TTS, widget", fill="#DCFCE7")
    rounded_box(d, boxes["llm"], "LLM Provider", "Gemini / Groq", fill="#FCE7F3")
    rounded_box(d, boxes["db"], "PostgreSQL", "User, quiz, attempt, session", fill="#E2E8F0")
    rounded_box(d, boxes["memory"], "mem0 Memory", "Lỗi sai, sở thích, tiến độ", fill="#EDE9FE")
    arrow(d, (380, 320), (540, 250))
    arrow(d, (715, 330), (715, 470), fill="#64748B")
    arrow(d, (890, 250), (1030, 250))
    arrow(d, (715, 470), (715, 390), fill="#64748B")
    arrow(d, (890, 550), (1030, 550), fill="#EA580C")
    arrow(d, (1205, 330), (1205, 470), fill="#7C3AED")
    arrow(d, (1380, 550), (1450, 550), fill="#7C3AED")
    arrow(d, (715, 630), (715, 770), fill="#334155")
    arrow(d, (1205, 630), (1205, 770), fill="#334155")
    d.text((90, 430), "Luồng thoại: WebRTC audio", font=font(25), fill="#1D4ED8")
    d.text((90, 470), "Luồng chữ/widget: LiveKit data channel", font=font(25), fill="#1D4ED8")
    d.text((90, 510), "Luồng quản trị: REST API + JWT", font=font(25), fill="#1D4ED8")
    img.save(path)


def make_use_case(path: Path):
    img = Image.new("RGB", (1800, 1100), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "Biểu đồ use case tổng quan", font=font(42, bold=True), fill="#111827")
    boundary = (420, 135, 1540, 1000)
    d.rounded_rectangle(boundary, radius=25, outline="#1E3A8A", width=4)
    d.text((820, 165), "EnglishAItutor System", font=font(28, bold=True), fill="#1E3A8A")

    def actor(x, y, label):
        d.ellipse((x, y, x + 70, y + 70), outline="#111827", width=4)
        d.line((x + 35, y + 70, x + 35, y + 190), fill="#111827", width=4)
        d.line((x - 20, y + 115, x + 90, y + 115), fill="#111827", width=4)
        d.line((x + 35, y + 190, x - 15, y + 275), fill="#111827", width=4)
        d.line((x + 35, y + 190, x + 85, y + 275), fill="#111827", width=4)
        d.text((x - 20, y + 290), label, font=font(24, bold=True), fill="#111827")

    actor(120, 245, "Học viên")
    actor(120, 665, "Quản trị viên")
    actor(1600, 430, "AI/LiveKit")

    use_cases = [
        (560, 230, "Đăng ký / đăng nhập"),
        (910, 230, "Xem lộ trình học"),
        (1240, 230, "Mua gói học"),
        (560, 425, "Gia sư AI voice/text"),
        (910, 425, "Làm quiz trong chat"),
        (1240, 425, "Nhận sửa lỗi & recap"),
        (560, 620, "Làm quiz theo bộ"),
        (910, 620, "Thi nâng cấp"),
        (1240, 620, "Xem AI review"),
        (560, 815, "Quản lý người dùng"),
        (910, 815, "Quản lý kho quiz"),
        (1240, 815, "Duyệt thanh toán QR"),
    ]
    centers = {}
    for x, y, text in use_cases:
        d.ellipse((x, y, x + 270, y + 95), fill="#F8FAFC", outline="#64748B", width=3)
        draw_wrapped(d, (x + 20, y + 28), text, 230, font(21), align="center")
        centers[text] = (x + 135, y + 48)
    for text in list(centers)[:9]:
        d.line((190, 360, centers[text][0] - 135, centers[text][1]), fill="#64748B", width=2)
    for text in list(centers)[9:]:
        d.line((190, 780, centers[text][0] - 135, centers[text][1]), fill="#64748B", width=2)
    for text in ["Gia sư AI voice/text", "Làm quiz trong chat", "Nhận sửa lỗi & recap"]:
        d.line((1600, 570, centers[text][0] + 135, centers[text][1]), fill="#64748B", width=2)
    img.save(path)


def make_vopc(path: Path, title: str, blocks: list[tuple[str, str, list[str]]]):
    img = Image.new("RGB", (1600, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((55, 40), title, font=font(40, bold=True), fill="#111827")
    positions = [(110, 250, 350, 390), (520, 160, 900, 450), (1080, 230, 1450, 500), (520, 620, 900, 880)]
    colors = ["#FEF9C3", "#DCFCE7", "#E0F2FE", "#FCE7F3"]
    for idx, (stereotype, name, items) in enumerate(blocks):
        x1, y1, x2, y2 = positions[idx]
        d.rounded_rectangle((x1, y1, x2, y2), radius=8, fill=colors[idx], outline="#991B1B", width=3)
        d.text((x1 + 18, y1 + 14), f"<<{stereotype}>>", font=font(21), fill="#111827")
        d.text((x1 + 18, y1 + 44), name, font=font(24, bold=True), fill="#111827")
        d.line((x1, y1 + 84, x2, y1 + 84), fill="#991B1B", width=2)
        y = y1 + 100
        item_limit = 2 if idx == 0 else 8
        for item in items[:item_limit]:
            d.text((x1 + 18, y), f"• {item}", font=font(16), fill="#111827")
            y += 24
    arrow(d, (350, 320), (520, 305), fill="#991B1B", width=3)
    arrow(d, (900, 305), (1080, 335), fill="#991B1B", width=3)
    arrow(d, (710, 450), (710, 620), fill="#991B1B", width=3)
    img.save(path)


def make_sequence(path: Path, title: str, actors: list[str], steps: list[tuple[int, int, str]]):
    img = Image.new("RGB", (1700, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((55, 40), title, font=font(40, bold=True), fill="#111827")
    xs = [170 + i * 300 for i in range(len(actors))]
    top = 150
    bottom = 910
    for x, actor_name in zip(xs, actors):
        d.rounded_rectangle((x - 110, top, x + 110, top + 70), radius=8, fill="#EFF6FF", outline="#1E40AF", width=3)
        draw_wrapped(d, (x - 95, top + 15), actor_name, 190, font(18, bold=True), align="center")
        d.line((x, top + 70, x, bottom), fill="#94A3B8", width=3)
    y = 260
    for i, j, msg in steps:
        x1, x2 = xs[i], xs[j]
        arrow(d, (x1, y), (x2, y), fill="#334155", width=3)
        label_x = min(x1, x2) + 10
        draw_wrapped(d, (label_x, y - 38), msg, abs(x2 - x1) - 20, font(17), fill="#111827")
        y += 82
    img.save(path)


def make_erd(path: Path):
    img = Image.new("RGB", (1850, 1150), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 45), "Thiết kế cơ sở dữ liệu chính", font=font(42, bold=True), fill="#111827")

    def entity(x, y, w, h, name, fields, fill):
        d.rounded_rectangle((x, y, x + w, y + h), radius=15, fill=fill, outline="#334155", width=3)
        d.text((x + 20, y + 16), name, font=font(26, bold=True), fill="#111827")
        yy = y + 58
        for field in fields:
            d.text((x + 20, yy), field, font=font(20), fill="#111827")
            yy += 29
        return (x, y, x + w, y + h)

    users = entity(80, 170, 360, 270, "users", ["id PK", "email unique", "password_hash", "name", "cefr_level", "role", "subscription_plan"], "#DBEAFE")
    refresh = entity(80, 530, 360, 210, "refresh_tokens", ["id PK", "user_id FK", "jti unique", "expires_at", "revoked_at"], "#FCE7F3")
    reset = entity(80, 820, 360, 210, "password_reset_codes", ["id PK", "user_id FK", "email", "code_hash", "expires_at", "used_at"], "#FEE2E2")
    sessions = entity(560, 170, 420, 310, "practice_sessions", ["id PK", "user_id FK", "room_name", "topic / level", "turns, errors", "grammar/vocab/fluency", "stats_json"], "#DCFCE7")
    errors = entity(1110, 190, 360, 230, "error_logs", ["id PK", "session_id FK", "user_id FK", "error_type", "original_text", "corrected_text"], "#FEF3C7")
    sets = entity(560, 600, 420, 250, "quiz_sets", ["id PK", "created_by FK", "title", "source / preset", "topic / level", "license / attribution"], "#EDE9FE")
    quizzes = entity(1110, 560, 360, 270, "quizzes", ["id PK", "user_id FK", "quiz_set_id FK", "title", "level/source", "questions_json"], "#E0F2FE")
    attempts = entity(1110, 870, 360, 230, "quiz_attempts", ["id PK", "quiz_id FK", "user_id FK", "answers_json", "ai_review_json", "score"], "#F0FDFA")
    payments = entity(1530, 560, 280, 230, "payment_requests", ["id PK", "user_id FK", "plan", "amount_vnd", "status", "qr_payload"], "#FFEDD5")
    arrow(d, (440, 300), (560, 310), fill="#15803D")
    d.text((465, 270), "1-N", font=font(19), fill="#111827")
    arrow(d, (440, 630), (560, 700), fill="#7C3AED")
    arrow(d, (440, 900), (560, 780), fill="#991B1B")
    arrow(d, (980, 325), (1110, 305), fill="#EA580C")
    d.text((1010, 285), "1-N", font=font(19), fill="#111827")
    arrow(d, (980, 725), (1110, 690), fill="#2563EB")
    arrow(d, (1290, 830), (1290, 870), fill="#0F766E")
    arrow(d, (440, 330), (1530, 650), fill="#EA580C", width=3)
    img.save(path)


def make_wireframe(path: Path, title: str, sections: list[str]):
    img = Image.new("RGB", (1600, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((55, 40), title, font=font(40, bold=True), fill="#111827")
    d.rounded_rectangle((170, 145, 1430, 860), radius=20, fill="#F8FAFC", outline="#94A3B8", width=3)
    d.rectangle((170, 145, 1430, 220), fill="#E0F2FE", outline="#94A3B8", width=3)
    d.text((205, 170), "EnglishAItutor", font=font(27, bold=True), fill="#0F172A")
    for i, sec in enumerate(sections):
        x = 235 + (i % 3) * 400
        y = 285 + (i // 3) * 230
        d.rounded_rectangle((x, y, x + 330, y + 150), radius=12, fill="white", outline="#CBD5E1", width=3)
        d.text((x + 20, y + 25), sec, font=font(24, bold=True), fill="#111827")
        d.line((x + 20, y + 75, x + 295, y + 75), fill="#CBD5E1", width=3)
        d.line((x + 20, y + 105, x + 250, y + 105), fill="#E2E8F0", width=3)
    img.save(path)


def generate_assets():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture(DIAGRAM_DIR / "architecture.png")
    make_use_case(DIAGRAM_DIR / "use_cases.png")
    make_vopc(
        DIAGRAM_DIR / "vopc_login.png",
        "VOPC use case Đăng nhập",
        [
            ("actor", "Học viên", ["nhập email", "nhập mật khẩu", "gửi yêu cầu"]),
            ("boundary", "LoginPage", ["hiển thị form", "validate dữ liệu", "lưu access token"]),
            ("control", "AuthRouter", ["kiểm tra mật khẩu", "tạo JWT", "tạo refresh cookie"]),
            ("entity", "User / RefreshToken", ["email", "password_hash", "role", "jti", "expires_at"]),
        ],
    )
    make_sequence(
        DIAGRAM_DIR / "seq_login.png",
        "Sơ đồ trình tự use case Đăng nhập",
        ["Học viên", "Frontend", "Auth API", "Database"],
        [
            (0, 1, "Nhập email, mật khẩu"),
            (1, 2, "POST /auth/login"),
            (2, 3, "Tìm user và refresh token"),
            (3, 2, "Trả thông tin user"),
            (2, 1, "Access JWT + refresh cookie"),
            (1, 0, "Điều hướng theo role"),
        ],
    )
    make_vopc(
        DIAGRAM_DIR / "vopc_practice.png",
        "VOPC use case Bắt đầu phiên Gia sư AI",
        [
            ("actor", "Học viên", ["chọn Gia sư AI", "bấm bắt đầu", "nói hoặc gõ"]),
            ("boundary", "PracticePage", ["xin token", "join LiveKit room", "render transcript/widget"]),
            ("control", "TokenRouter / Agent", ["kiểm tra quota", "dispatch room", "STT-LangGraph-TTS"]),
            ("entity", "PracticeSession", ["room_name", "topic", "level", "turns", "scores", "stats_json"]),
        ],
    )
    make_sequence(
        DIAGRAM_DIR / "seq_practice.png",
        "Sơ đồ trình tự phiên Gia sư AI",
        ["Học viên", "Frontend", "Token API", "LiveKit", "Agent"],
        [
            (0, 1, "Bấm bắt đầu phiên"),
            (1, 2, "POST /token kèm level/topic"),
            (2, 1, "Trả roomName và JWT LiveKit"),
            (1, 3, "Join room, publish audio/data"),
            (3, 4, "Dispatch agent theo room metadata"),
            (4, 3, "Trả audio TTS, transcript, widget"),
            (3, 1, "Data/audio về trình duyệt"),
        ],
    )
    make_vopc(
        DIAGRAM_DIR / "vopc_chat_quiz.png",
        "VOPC use case Làm quiz trong khung chat",
        [
            ("actor", "Học viên", ["yêu cầu làm bài", "chọn đáp án", "nộp bài"]),
            ("boundary", "InlineQuizWidget", ["hiển thị câu hỏi", "ghi đáp án", "hiển thị review"]),
            ("control", "Tutor Agent", ["nhận diện intent", "chọn quiz bank", "emit widget"]),
            ("entity", "Quiz / QuizAttempt", ["questions_json", "answers_json", "score", "ai_review_json"]),
        ],
    )
    make_sequence(
        DIAGRAM_DIR / "seq_chat_quiz.png",
        "Sơ đồ trình tự quiz trong chat",
        ["Học viên", "Frontend", "LiveKit", "Tutor Agent", "Quiz DB"],
        [
            (0, 1, "Gõ: I want to do exercises"),
            (1, 2, "Gửi data channel"),
            (2, 3, "Agent nhận text"),
            (3, 4, "Lấy câu phù hợp trình độ/chủ đề"),
            (4, 3, "Trả bộ câu trắc nghiệm"),
            (3, 2, "Emit ai-quiz-widget"),
            (2, 1, "Render widget trong chat"),
        ],
    )
    make_vopc(
        DIAGRAM_DIR / "vopc_admin_quiz.png",
        "VOPC use case Quản lý quiz",
        [
            ("actor", "Quản trị viên", ["mở kho quiz", "tạo/sửa/xoá", "import nguồn"]),
            ("boundary", "QuizStudio", ["form quiz", "upload file", "danh sách bộ đề"]),
            ("control", "QuizRouter", ["validate role", "parse dữ liệu", "gọi AI sinh câu"]),
            ("entity", "QuizSet / Quiz", ["source", "level", "questions_json", "license"]),
        ],
    )
    make_sequence(
        DIAGRAM_DIR / "seq_admin_quiz.png",
        "Sơ đồ trình tự quản lý quiz",
        ["Admin", "QuizStudio", "Quiz API", "LLM", "Database"],
        [
            (0, 1, "Nhập nội dung hoặc chọn import"),
            (1, 2, "POST /quizzes/admin/import-source"),
            (2, 3, "Sinh câu hỏi theo nguồn/trình độ"),
            (3, 2, "Trả JSON câu hỏi"),
            (2, 4, "Lưu quiz_set và quizzes"),
            (4, 2, "Xác nhận lưu thành công"),
            (2, 1, "Hiển thị bộ quiz mới"),
        ],
    )
    make_vopc(
        DIAGRAM_DIR / "vopc_payment.png",
        "VOPC use case Thanh toán QR",
        [
            ("actor", "Học viên", ["chọn gói", "quét QR", "chờ duyệt"]),
            ("boundary", "BillingPage", ["hiển thị gói", "tạo QR", "trạng thái"]),
            ("control", "BillingRouter", ["tính giá", "tạo request", "kiểm tra quota"]),
            ("entity", "PaymentRequest / User", ["plan", "amount_vnd", "status", "subscription_plan"]),
        ],
    )
    make_sequence(
        DIAGRAM_DIR / "seq_payment.png",
        "Sơ đồ trình tự thanh toán QR",
        ["Học viên", "BillingPage", "Billing API", "Admin", "Database"],
        [
            (0, 1, "Chọn gói Plus/Ultra"),
            (1, 2, "POST /billing/payment-requests"),
            (2, 4, "Lưu payment_request pending"),
            (2, 1, "Trả QR payload"),
            (3, 2, "Duyệt thanh toán"),
            (2, 4, "Cập nhật status và plan user"),
            (2, 3, "Trả kết quả duyệt"),
        ],
    )
    make_erd(DIAGRAM_DIR / "erd.png")
    make_wireframe(DIAGRAM_DIR / "wire_landing.png", "Thiết kế giao diện Landing page", ["Hero giới thiệu", "Gói học", "Lợi ích", "CTA đăng ký"])
    make_wireframe(DIAGRAM_DIR / "wire_practice.png", "Thiết kế giao diện Gia sư AI", ["Header phiên", "Khung chat", "Widget quiz", "Micro & nhập text", "Recap lỗi sai"])
    make_wireframe(DIAGRAM_DIR / "wire_quiz.png", "Thiết kế giao diện Kho quiz", ["Danh sách bộ đề", "Lọc theo trình độ", "Màn làm bài", "AI review"])
    make_wireframe(DIAGRAM_DIR / "wire_admin.png", "Thiết kế giao diện Quản trị", ["Quản lý quiz", "Import nguồn", "Người dùng", "Thanh toán QR"])


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 90, 110, 90, 110)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
            if header and r_idx == 0:
                set_cell_shading(cell, "EAF2F8")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_bookmark(paragraph, name: str, bm_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, anchor: str, text: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instr: str, placeholder: str = ""):
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(begin)
    r._r.append(instr_el)
    r._r.append(separate)
    r._r.append(text)
    r._r.append(end)


def set_update_fields(docx_path: Path):
    import shutil
    import tempfile
    from lxml import etree

    tmp = Path(tempfile.mkdtemp(prefix="aitutor_docx_"))
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            z.extractall(tmp)
        settings = tmp / "word" / "settings.xml"
        parser = etree.XMLParser(remove_blank_text=False)
        tree = etree.parse(str(settings), parser)
        root = tree.getroot()
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        uf = root.find(f"{{{ns}}}updateFields")
        if uf is None:
            uf = etree.Element(f"{{{ns}}}updateFields")
            root.insert(0, uf)
        uf.set(f"{{{ns}}}val", "true")
        tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
        with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for p in tmp.rglob("*"):
                if p.is_file():
                    z.write(p, p.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


HEADING_OUTLINE = [
    (1, "MỞ ĐẦU", "h_open"),
    (2, "1. Lý do chọn đề tài", "h_open_reason"),
    (2, "2. Mục đích nghiên cứu", "h_open_goal"),
    (2, "3. Đối tượng và phạm vi nghiên cứu", "h_open_scope"),
    (2, "4. Phương pháp nghiên cứu", "h_open_method"),
    (2, "5. Kết quả dự kiến", "h_open_result"),
    (2, "6. Bố cục báo cáo", "h_open_structure"),
    (1, "CHƯƠNG 1. CƠ SỞ LÝ THUYẾT", "h_ch1"),
    (2, "1.1. Tổng quan về hệ thống gia sư AI", "h_1_1"),
    (2, "1.2. React, TypeScript và Vite", "h_1_2"),
    (2, "1.3. FastAPI và kiến trúc REST", "h_1_3"),
    (2, "1.4. WebRTC và LiveKit", "h_1_4"),
    (2, "1.5. Mô hình ngôn ngữ lớn và LangGraph", "h_1_5"),
    (2, "1.6. Nhận dạng giọng nói, tổng hợp giọng nói và VAD", "h_1_6"),
    (2, "1.7. PostgreSQL, pgvector và bộ nhớ dài hạn", "h_1_7"),
    (1, "CHƯƠNG 2. KHẢO SÁT, PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "h_ch2"),
    (2, "2.1. Khảo sát hệ thống", "h_2_1"),
    (3, "2.1.1. Giới thiệu hệ thống", "h_2_1_1"),
    (3, "2.1.2. Các yêu cầu của hệ thống", "h_2_1_2"),
    (2, "2.2. Phân tích hệ thống", "h_2_2"),
    (3, "2.2.1. Mô hình use case", "h_2_2_1"),
    (3, "2.2.2. Đặc tả các use case", "h_2_2_2"),
    (3, "2.2.3. Phân tích use case tiêu biểu", "h_2_2_3"),
    (2, "2.3. Thiết kế hệ thống", "h_2_3"),
    (3, "2.3.1. Thiết kế kiến trúc", "h_2_3_1"),
    (3, "2.3.2. Thiết kế cơ sở dữ liệu", "h_2_3_2"),
    (3, "2.3.3. Thiết kế giao diện", "h_2_3_3"),
    (1, "CHƯƠNG 3. CÀI ĐẶT VÀ KẾT QUẢ", "h_ch3"),
    (2, "3.1. Cài đặt hệ thống", "h_3_1"),
    (3, "3.1.1. Yêu cầu cấu hình", "h_3_1_1"),
    (3, "3.1.2. Hướng dẫn triển khai", "h_3_1_2"),
    (2, "3.2. Kết quả thực hiện", "h_3_2"),
    (3, "3.2.1. Giao diện người dùng", "h_3_2_1"),
    (3, "3.2.2. Giao diện quản trị", "h_3_2_2"),
    (3, "3.2.3. Kết quả xử lý nghiệp vụ AI", "h_3_2_3"),
    (2, "3.3. Kiểm thử", "h_3_3"),
    (3, "3.3.1. Kế hoạch kiểm thử", "h_3_3_1"),
    (3, "3.3.2. Lịch trình kiểm thử", "h_3_3_2"),
    (3, "3.3.3. Kết quả kiểm thử", "h_3_3_3"),
    (3, "3.3.4. Điều kiện chấp nhận kiểm thử", "h_3_3_4"),
    (1, "KẾT LUẬN", "h_conclusion"),
    (1, "TÀI LIỆU THAM KHẢO", "h_refs"),
]

FIGURES = [
    ("fig_1_1", "Hình 1.1. Kiến trúc tổng thể hệ thống EnglishAItutor"),
    ("fig_2_1", "Hình 2.1. Biểu đồ use case tổng quan của hệ thống"),
    ("fig_2_2", "Hình 2.2. VOPC use case Đăng nhập"),
    ("fig_2_3", "Hình 2.3. Sơ đồ trình tự use case Đăng nhập"),
    ("fig_2_4", "Hình 2.4. VOPC use case Bắt đầu phiên Gia sư AI"),
    ("fig_2_5", "Hình 2.5. Sơ đồ trình tự phiên Gia sư AI"),
    ("fig_2_6", "Hình 2.6. VOPC use case Làm quiz trong khung chat"),
    ("fig_2_7", "Hình 2.7. Sơ đồ trình tự quiz trong chat"),
    ("fig_2_8", "Hình 2.8. VOPC use case Quản lý quiz"),
    ("fig_2_9", "Hình 2.9. Sơ đồ trình tự quản lý quiz"),
    ("fig_2_10", "Hình 2.10. VOPC use case Thanh toán QR"),
    ("fig_2_11", "Hình 2.11. Sơ đồ trình tự thanh toán QR"),
    ("fig_2_12", "Hình 2.12. Thiết kế cơ sở dữ liệu chính"),
    ("fig_2_13", "Hình 2.13. Thiết kế giao diện Landing page"),
    ("fig_2_14", "Hình 2.14. Thiết kế giao diện Gia sư AI"),
    ("fig_2_15", "Hình 2.15. Thiết kế giao diện Kho quiz"),
    ("fig_2_16", "Hình 2.16. Thiết kế giao diện Quản trị"),
    ("fig_3_1", "Hình 3.1. Giao diện Landing page"),
    ("fig_3_2", "Hình 3.2. Giao diện đăng nhập"),
    ("fig_3_3", "Hình 3.3. Giao diện trang học tập"),
    ("fig_3_4", "Hình 3.4. Giao diện Gia sư AI"),
    ("fig_3_5", "Hình 3.5. Giao diện kho quiz"),
    ("fig_3_6", "Hình 3.6. Giao diện thanh toán QR"),
]

TABLES = [
    ("tbl_2_1", "Bảng 2.1. Tác nhân của hệ thống"),
    ("tbl_2_2", "Bảng 2.2. Yêu cầu chức năng"),
    ("tbl_2_3", "Bảng 2.3. Yêu cầu phi chức năng"),
    ("tbl_2_4", "Bảng 2.4. Đặc tả use case Đăng nhập"),
    ("tbl_2_5", "Bảng 2.5. Đặc tả use case Bắt đầu phiên Gia sư AI"),
    ("tbl_2_6", "Bảng 2.6. Đặc tả use case Làm quiz trong chat"),
    ("tbl_2_7", "Bảng 2.7. Đặc tả use case Làm quiz theo bộ"),
    ("tbl_2_8", "Bảng 2.8. Đặc tả use case Quản lý quiz"),
    ("tbl_2_9", "Bảng 2.9. Đặc tả use case Thanh toán QR"),
    ("tbl_2_10", "Bảng 2.10. Đặc tả use case AI review"),
    ("tbl_2_11", "Bảng 2.11. Đặc tả use case Thi nâng cấp"),
    ("tbl_2_12", "Bảng 2.12. Đặc tả use case Quản lý người dùng"),
    ("tbl_2_13", "Bảng 2.13. Mô tả bảng users"),
    ("tbl_2_14", "Bảng 2.14. Mô tả bảng practice_sessions"),
    ("tbl_2_15", "Bảng 2.15. Mô tả bảng error_logs"),
    ("tbl_2_16", "Bảng 2.16. Mô tả bảng refresh_tokens"),
    ("tbl_2_17", "Bảng 2.17. Mô tả bảng quiz_sets"),
    ("tbl_2_18", "Bảng 2.18. Mô tả bảng quizzes"),
    ("tbl_2_19", "Bảng 2.19. Mô tả bảng quiz_attempts"),
    ("tbl_2_20", "Bảng 2.20. Mô tả bảng payment_requests"),
    ("tbl_3_1", "Bảng 3.1. Thành phần triển khai hệ thống"),
    ("tbl_3_2", "Bảng 3.2. Cấu hình môi trường"),
    ("tbl_3_3", "Bảng 3.3. Các API chính của hệ thống"),
    ("tbl_3_4", "Bảng 3.4. Luồng xử lý AI trong phiên Gia sư"),
    ("tbl_3_5", "Bảng 3.5. Lịch trình kiểm thử"),
    ("tbl_3_6", "Bảng 3.6. Kết quả kiểm thử chức năng"),
]


class Builder:
    def __init__(self):
        self.doc = Document()
        self.bm_id = 1
        self.figure_map = {caption: bid for bid, caption in FIGURES}
        self.table_map = {caption: bid for bid, caption in TABLES}
        self.setup_doc()

    def setup_doc(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        normal = self.doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        normal.font.size = Pt(13)
        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = self.doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
        self.doc.styles["Heading 1"].font.size = Pt(16)
        self.doc.styles["Heading 2"].font.size = Pt(14)
        self.doc.styles["Heading 3"].font.size = Pt(13)

    def bookmark(self, paragraph, name):
        add_bookmark(paragraph, name, self.bm_id)
        self.bm_id += 1

    def heading(self, text: str, level: int, bookmark: str):
        p = self.doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.bookmark(p, bookmark)
        return p

    def para(self, text: str = "", align=None, bold=False, italic=False, first_line=True):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        if first_line:
            p.paragraph_format.first_line_indent = Cm(1.0)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        return p

    def bullets(self, items):
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.add_run(item)

    def caption(self, text: str, kind: str = "fig"):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        bid = self.figure_map.get(text) if kind == "fig" else self.table_map.get(text)
        if bid:
            self.bookmark(p, bid)

    def add_table_caption(self, text: str):
        self.caption(text, kind="tbl")

    def image(self, path: Path, width_inches: float, caption: str):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
        self.caption(caption, "fig")

    def table(self, rows, caption: str, widths: list[float] | None = None):
        self.add_table_caption(caption)
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        style_table(table)
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or j in {0, 2, 3}) else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(value))
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                if i == 0:
                    run.bold = True
            if widths:
                for j, w in enumerate(widths):
                    table.cell(i, j).width = Cm(w)
        self.para("", first_line=False)
        return table

    def usecase_table(self, caption, name, desc, actor, trigger, pre, post, main, alt):
        rows = [
            ["Tên usecase", name],
            ["Mô tả", desc],
            ["Actor", actor],
            ["Điều kiện kích hoạt", trigger],
            ["Tiền điều kiện", pre],
            ["Hậu điều kiện", post],
            ["Luồng cơ bản", "\n".join(f"{i + 1}. {s}" for i, s in enumerate(main))],
            ["Luồng rẽ nhánh", "\n".join(alt) if alt else "Không"],
        ]
        self.add_table_caption(caption)
        table = self.doc.add_table(rows=len(rows), cols=2)
        style_table(table, header=False)
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = ""
                if j == 0:
                    set_cell_shading(cell, "F8FAFC")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(value)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                if j == 0:
                    run.bold = True
        self.para("", first_line=False)

    def cover(self):
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Cm(1.7)
        table.columns[1].width = Cm(15.0)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        for cell in (left, right):
            set_cell_margins(cell, 80, 80, 80, 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tc_pr = left._tc.get_or_add_tcPr()
        text_dir = OxmlElement("w:textDirection")
        text_dir.set(qn("w:val"), "btLr")
        tc_pr.append(text_dir)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("NGUYỄN VIỆT ANH\nHỆ THỐNG THÔNG TIN")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for text, size, bold in [
            ("BỘ CÔNG THƯƠNG", 16, True),
            ("TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP HÀ NỘI", 16, True),
            ("--------------------------------", 14, False),
            ("\n\nĐỒ ÁN TỐT NGHIỆP", 26, True),
            ("NGÀNH HỆ THỐNG THÔNG TIN", 18, False),
            ("\n\nXÂY DỰNG HỆ THỐNG GIA SƯ AI TIẾNG ANH\nTÍCH HỢP LUYỆN NÓI, QUIZ VÀ ĐÁNH GIÁ CÁ NHÂN HÓA", 18, True),
            ("\n\n\nCBHD: TS. Giang Thành Trung", 15, True),
            ("Sinh viên: Nguyễn Việt Anh", 15, True),
            ("Mã sinh viên: 2021600938", 15, True),
            ("\n\n\nHà Nội - 2026", 16, True),
        ]:
            run = rp.add_run(text + "\n")
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
        self.doc.add_page_break()

    def front_matter(self):
        self.heading("LỜI CẢM ƠN", 1, "ack")
        self.para(
            "Em xin gửi lời cảm ơn chân thành tới các thầy cô khoa Công nghệ thông tin, Trường Đại học Công nghiệp Hà Nội "
            "đã trang bị cho em nền tảng kiến thức trong quá trình học tập. Em đặc biệt cảm ơn giảng viên hướng dẫn đã góp ý, "
            "định hướng và hỗ trợ em trong quá trình thực hiện đồ án tốt nghiệp."
        )
        self.para(
            "Đồ án được xây dựng với mục tiêu tạo ra một hệ thống học tiếng Anh có tính ứng dụng thực tế: học viên có thể luyện nói "
            "với gia sư AI, làm bài quiz theo trình độ, nhận sửa lỗi và theo dõi tiến bộ. Trong quá trình thực hiện, em đã vận dụng "
            "các kiến thức về phát triển phần mềm, cơ sở dữ liệu, trí tuệ nhân tạo và triển khai hệ thống web."
        )
        self.para(
            "Do thời gian và kinh nghiệm còn hạn chế, báo cáo khó tránh khỏi thiếu sót. Em mong nhận được ý kiến đóng góp của thầy cô "
            "để sản phẩm và báo cáo được hoàn thiện hơn.",
        )
        self.doc.add_page_break()
        self.heading("MỤC LỤC", 1, "toc")
        for level, title, bid in HEADING_OUTLINE:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(max(0, level - 1) * 0.65)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            add_internal_link(p, bid, title)
            p.add_run("\t")
            add_field(p, f"PAGEREF {bid} \\h", "")
        self.doc.add_page_break()
        self.heading("DANH MỤC HÌNH ẢNH", 1, "fig_list")
        for bid, caption in FIGURES:
            p = self.doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            add_internal_link(p, bid, caption)
            p.add_run("\t")
            add_field(p, f"PAGEREF {bid} \\h", "")
        self.doc.add_page_break()
        self.heading("DANH MỤC BẢNG BIỂU", 1, "tbl_list")
        for bid, caption in TABLES:
            p = self.doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            add_internal_link(p, bid, caption)
            p.add_run("\t")
            add_field(p, f"PAGEREF {bid} \\h", "")
        self.doc.add_page_break()

    def intro(self):
        self.heading("MỞ ĐẦU", 1, "h_open")
        self.heading("1. Lý do chọn đề tài", 2, "h_open_reason")
        self.para(
            "Nhu cầu học tiếng Anh tại Việt Nam ngày càng tăng, đặc biệt ở nhóm học sinh, sinh viên cần luyện giao tiếp và ôn tập "
            "theo trình độ. Tuy nhiên, nhiều nền tảng học ngoại ngữ hiện nay hoặc chỉ dừng ở kho bài tập tĩnh, hoặc chỉ cung cấp "
            "chatbot hội thoại chưa gắn với lộ trình, lỗi sai và bài luyện tập cụ thể."
        )
        self.para(
            "Đề tài xây dựng hệ thống Gia sư AI tiếng Anh nhằm kết hợp luyện nói thời gian thực, sửa lỗi, tạo quiz và đánh giá kết quả "
            "trên cùng một sản phẩm. Điểm khác biệt của hệ thống là học viên không chỉ trò chuyện với AI mà còn có thể làm quiz ngay "
            "trong khung chat, làm bài theo bộ đề, nhận AI review và được giới hạn nội dung theo trình độ CEFR."
        )
        self.heading("2. Mục đích nghiên cứu", 2, "h_open_goal")
        self.bullets([
            "Xây dựng ứng dụng web học tiếng Anh với hai vai trò: học viên và quản trị viên.",
            "Tích hợp phiên Gia sư AI trực tuyến có giọng nói, văn bản, transcript và widget bài tập.",
            "Tổ chức kho quiz theo bộ đề, trình độ CEFR, nguồn dữ liệu và kết quả làm bài.",
            "Cung cấp cơ chế AI review để phân tích điểm mạnh, điểm yếu sau khi học viên làm bài.",
            "Xây dựng quản trị người dùng, quiz và thanh toán QR cho mô hình gói học.",
        ])
        self.heading("3. Đối tượng và phạm vi nghiên cứu", 2, "h_open_scope")
        self.para(
            "Đối tượng nghiên cứu là hệ thống học tiếng Anh trực tuyến ứng dụng AI cho học viên ở các trình độ A1-C2. Phạm vi triển khai "
            "tập trung vào sản phẩm MVP có thể sử dụng được: đăng ký/đăng nhập, phân quyền, Gia sư AI, quiz, thi nâng cấp, thanh toán QR "
            "và quản trị hệ thống."
        )
        self.heading("4. Phương pháp nghiên cứu", 2, "h_open_method")
        self.bullets([
            "Khảo sát luồng học tiếng Anh phổ biến: luyện nói, ôn lỗi, làm bài trắc nghiệm và đánh giá sau bài làm.",
            "Phân tích yêu cầu bằng use case, đặc tả nghiệp vụ và thiết kế cơ sở dữ liệu.",
            "Xây dựng hệ thống theo kiến trúc frontend React, backend FastAPI, LiveKit và agent AI.",
            "Kiểm thử chức năng theo các kịch bản chính của học viên và quản trị viên.",
        ])
        self.heading("5. Kết quả dự kiến", 2, "h_open_result")
        self.para(
            "Sản phẩm hoàn thiện ở mức MVP có thể trình diễn và sử dụng thử: học viên đăng nhập, luyện nói với AI, nhận giọng nói phản hồi, "
            "làm quiz trong chat hoặc trong kho bài, xem kết quả AI review; quản trị viên quản lý người dùng, bộ quiz và thanh toán."
        )
        self.heading("6. Bố cục báo cáo", 2, "h_open_structure")
        self.para(
            "Báo cáo gồm ba chương chính: Chương 1 trình bày cơ sở lý thuyết; Chương 2 khảo sát, phân tích và thiết kế hệ thống; "
            "Chương 3 mô tả cài đặt, kết quả thực hiện và kiểm thử. Cuối báo cáo là phần kết luận và tài liệu tham khảo."
        )
        self.doc.add_page_break()

    def chapter1(self):
        self.heading("CHƯƠNG 1. CƠ SỞ LÝ THUYẾT", 1, "h_ch1")
        self.heading("1.1. Tổng quan về hệ thống gia sư AI", 2, "h_1_1")
        self.para(
            "Gia sư AI là mô hình ứng dụng trí tuệ nhân tạo vào quá trình học tập cá nhân hóa. Hệ thống không chỉ trả lời câu hỏi như chatbot "
            "thông thường mà còn đóng vai trò người hướng dẫn: đặt câu hỏi, điều chỉnh độ khó, sửa lỗi, tạo bài tập và ghi nhận tiến độ học viên."
        )
        self.image(DIAGRAM_DIR / "architecture.png", 6.5, "Hình 1.1. Kiến trúc tổng thể hệ thống EnglishAItutor")
        self.heading("1.2. React, TypeScript và Vite", 2, "h_1_2")
        self.para(
            "React được sử dụng để xây dựng giao diện dạng component, giúp tách các màn hình học viên, quản trị viên và widget chat thành các phần "
            "dễ bảo trì. TypeScript bổ sung kiểm tra kiểu tĩnh, giảm lỗi khi truyền dữ liệu giữa API, store và component. Vite giúp quá trình phát triển "
            "frontend nhanh, hỗ trợ hot reload và build sản phẩm tối ưu."
        )
        self.heading("1.3. FastAPI và kiến trúc REST", 2, "h_1_3")
        self.para(
            "FastAPI là framework backend Python hỗ trợ khai báo API bằng type hint, tự sinh tài liệu OpenAPI và xử lý bất đồng bộ. Trong đề tài, "
            "FastAPI đảm nhiệm xác thực, quản lý người dùng, quiz, thanh toán, token LiveKit và lưu trữ dữ liệu phiên học."
        )
        self.heading("1.4. WebRTC và LiveKit", 2, "h_1_4")
        self.para(
            "WebRTC cho phép trình duyệt truyền âm thanh thời gian thực mà không cần plugin. LiveKit được sử dụng như lớp hạ tầng phòng học trực tuyến: "
            "frontend tham gia phòng, agent AI được dispatch vào cùng phòng, hai bên trao đổi audio và data channel. Data channel cũng được tận dụng để gửi "
            "transcript, tin nhắn chữ và widget quiz."
        )
        self.heading("1.5. Mô hình ngôn ngữ lớn và LangGraph", 2, "h_1_5")
        self.para(
            "Mô hình ngôn ngữ lớn được dùng để tạo phản hồi hội thoại, phân tích lỗi và sinh nhận xét sau bài làm. LangGraph giúp tổ chức luồng xử lý agent "
            "thành các node rõ ràng như phân tích lỗi, định tuyến phản hồi, sửa lỗi, đổi chủ đề, lưu memory và tạo bài tập."
        )
        self.heading("1.6. Nhận dạng giọng nói, tổng hợp giọng nói và VAD", 2, "h_1_6")
        self.para(
            "STT chuyển giọng nói của học viên thành văn bản để agent phân tích. TTS chuyển phản hồi của agent thành âm thanh để học viên nghe được. "
            "VAD phát hiện vùng có tiếng nói, giúp hệ thống tránh ngắt câu quá sớm và giảm hiện tượng agent chen ngang khi người học chưa nói xong."
        )
        self.heading("1.7. PostgreSQL, pgvector và bộ nhớ dài hạn", 2, "h_1_7")
        self.para(
            "PostgreSQL lưu dữ liệu nghiệp vụ như tài khoản, phiên học, lỗi sai, quiz, kết quả làm bài và thanh toán. Bên cạnh dữ liệu quan hệ, hệ thống có "
            "bộ nhớ dài hạn qua mem0/pgvector để lưu các thông tin có ích cho cá nhân hóa như lỗi thường gặp, sở thích và bằng chứng trình độ."
        )
        self.doc.add_page_break()

    def chapter2(self):
        self.heading("CHƯƠNG 2. KHẢO SÁT, PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1, "h_ch2")
        self.heading("2.1. Khảo sát hệ thống", 2, "h_2_1")
        self.heading("2.1.1. Giới thiệu hệ thống", 3, "h_2_1_1")
        self.para(
            "EnglishAItutor là hệ thống web hỗ trợ học tiếng Anh theo hướng all-in-one. Học viên có thể vào landing page, tạo tài khoản, mua gói học, "
            "luyện nói với gia sư AI, làm quiz theo bộ đề hoặc quiz được mở trực tiếp trong chat, sau đó nhận đánh giá điểm mạnh và điểm yếu."
        )
        self.heading("2.1.2. Các yêu cầu của hệ thống", 3, "h_2_1_2")
        self.table(
            [
                ["Tác nhân", "Mô tả", "Quyền chính"],
                ["Học viên", "Người sử dụng sản phẩm để học tiếng Anh", "Luyện nói, làm quiz, xem kết quả, mua gói"],
                ["Quản trị viên", "Người vận hành hệ thống", "Quản lý người dùng, quiz, nguồn import, thanh toán"],
                ["Gia sư AI", "Agent tham gia phòng LiveKit", "Hội thoại, sửa lỗi, tạo widget bài tập, TTS/STT"],
            ],
            "Bảng 2.1. Tác nhân của hệ thống",
            [3.0, 6.0, 6.0],
        )
        self.table(
            [
                ["Mã", "Yêu cầu chức năng", "Mô tả"],
                ["FR01", "Xác thực và phân quyền", "Đăng ký, đăng nhập, refresh token, quên mật khẩu, role learner/admin"],
                ["FR02", "Gia sư AI", "Học viên luyện nói hoặc nhắn tin với agent qua LiveKit"],
                ["FR03", "Sửa lỗi và cá nhân hóa", "Phân tích lỗi, lưu lỗi sai, dùng memory để điều chỉnh phản hồi"],
                ["FR04", "Quiz theo bộ", "Học viên làm quiz theo trình độ; admin tạo, sửa, xoá, import quiz"],
                ["FR05", "Quiz trong chat", "Agent nhận diện nhu cầu làm bài và gửi widget trắc nghiệm trong khung chat"],
                ["FR06", "AI review", "Tổng hợp kết quả sau bài làm, nêu điểm mạnh, điểm yếu và hướng cải thiện"],
                ["FR07", "Thi nâng cấp", "Đánh giá điều kiện nâng cấp trình độ theo bài kiểm tra"],
                ["FR08", "Gói học và QR", "Giới hạn lượt học theo gói, tạo yêu cầu thanh toán QR"],
            ],
            "Bảng 2.2. Yêu cầu chức năng",
            [2.0, 4.5, 8.5],
        )
        self.table(
            [
                ["Mã", "Yêu cầu phi chức năng", "Mô tả"],
                ["NFR01", "Bảo mật", "Mật khẩu hash, JWT ngắn hạn, refresh cookie, phân quyền API"],
                ["NFR02", "Tính thời gian thực", "Âm thanh và transcript phản hồi qua WebRTC/data channel"],
                ["NFR03", "Khả dụng", "Các service chạy bằng Docker Compose, có health check"],
                ["NFR04", "Dễ sử dụng", "Giao diện tiếng Việt, tách luồng học viên và quản trị viên"],
                ["NFR05", "Mở rộng", "Có thể bổ sung nguồn quiz, model AI, gói thanh toán và dạng bài mới"],
            ],
            "Bảng 2.3. Yêu cầu phi chức năng",
            [2.0, 4.5, 8.5],
        )

        self.heading("2.2. Phân tích hệ thống", 2, "h_2_2")
        self.heading("2.2.1. Mô hình use case", 3, "h_2_2_1")
        self.image(DIAGRAM_DIR / "use_cases.png", 6.5, "Hình 2.1. Biểu đồ use case tổng quan của hệ thống")
        self.heading("2.2.2. Đặc tả các use case", 3, "h_2_2_2")
        self.usecase_table(
            "Bảng 2.4. Đặc tả use case Đăng nhập",
            "Đăng nhập",
            "Cho phép người dùng xác thực để truy cập đúng giao diện theo vai trò.",
            "Học viên, quản trị viên",
            "Người dùng chọn chức năng đăng nhập",
            "Người dùng đã có tài khoản",
            "Hệ thống cấp access token và refresh token hợp lệ",
            ["Hiển thị form đăng nhập", "Người dùng nhập email và mật khẩu", "Hệ thống kiểm tra thông tin", "Sinh access JWT và refresh cookie", "Điều hướng theo role"],
            ["Sai mật khẩu: hiển thị thông báo lỗi", "Tài khoản không tồn tại: yêu cầu kiểm tra lại email"],
        )
        self.usecase_table(
            "Bảng 2.5. Đặc tả use case Bắt đầu phiên Gia sư AI",
            "Bắt đầu phiên Gia sư AI",
            "Học viên mở phòng LiveKit để luyện nói hoặc nhắn tin với agent.",
            "Học viên, Gia sư AI",
            "Học viên bấm nút bắt đầu phiên học",
            "Học viên đã đăng nhập và còn lượt chat",
            "Phòng LiveKit được tạo, agent tham gia và phiên học được lưu",
            ["Frontend gọi API lấy token", "Backend kiểm tra quota", "LiveKit tạo room và dispatch agent", "Agent chào học viên", "Học viên nói hoặc gõ để luyện tập"],
            ["Hết lượt chat: hiển thị widget đề nghị mua gói", "Không cấp quyền micro: hiển thị lỗi thiết bị"],
        )
        self.usecase_table(
            "Bảng 2.6. Đặc tả use case Làm quiz trong chat",
            "Làm quiz trong khung chat",
            "Agent gửi một bộ câu hỏi trắc nghiệm ngay trong khung trò chuyện khi học viên muốn luyện tập.",
            "Học viên, Gia sư AI",
            "Học viên nói/gõ muốn làm bài tập hoặc agent chủ động gợi ý sau một số lượt",
            "Phiên Gia sư AI đang hoạt động",
            "Widget quiz được hiển thị và học viên nhận nhận xét sau khi nộp",
            ["Agent nhận diện intent làm bài", "Agent chọn bài theo trình độ/chủ đề", "Frontend render widget", "Học viên chọn đáp án", "Hệ thống chấm và hiển thị review"],
            ["Không đủ dữ liệu quiz: agent tạo bài ngắn theo prompt", "Học viên từ chối: tiếp tục hội thoại"],
        )
        self.usecase_table(
            "Bảng 2.7. Đặc tả use case Làm quiz theo bộ",
            "Làm quiz theo bộ đề",
            "Học viên chọn một bộ quiz trong kho, trả lời từng câu và nộp bài.",
            "Học viên",
            "Học viên chọn mục Bài quiz",
            "Học viên đã đăng nhập và bài không bị khoá theo khoảng trình độ",
            "Kết quả, điểm số và AI review được lưu",
            ["Hiển thị danh sách bộ đề", "Lọc theo trình độ học viên", "Học viên chọn quiz", "Trả lời câu hỏi", "Nộp bài", "Hiển thị điểm và review"],
            ["Bài vượt quá trình độ cho phép: hiển thị trạng thái khoá"],
        )
        self.usecase_table(
            "Bảng 2.8. Đặc tả use case Quản lý quiz",
            "Quản lý quiz",
            "Quản trị viên tạo, sửa, xoá, import quiz từ file hoặc nguồn tham khảo.",
            "Quản trị viên",
            "Quản trị viên mở màn Quản lý quiz",
            "Người dùng có role admin",
            "Kho quiz được cập nhật",
            ["Xem danh sách quiz", "Tạo hoặc import quiz", "Sửa nội dung câu hỏi", "Xoá quiz không dùng", "Tạo bộ quiz từ nguồn mở hoặc sách"],
            ["File import sai định dạng: hệ thống trả lỗi", "Nguồn URL không truy cập được: yêu cầu thử nguồn khác"],
        )
        self.usecase_table(
            "Bảng 2.9. Đặc tả use case Thanh toán QR",
            "Thanh toán QR",
            "Học viên chọn gói học và tạo yêu cầu thanh toán bằng QR.",
            "Học viên, quản trị viên",
            "Học viên chọn gói Plus hoặc Ultra",
            "Học viên đã đăng nhập",
            "Yêu cầu thanh toán được tạo và chờ quản trị viên duyệt",
            ["Học viên chọn gói", "Hệ thống tạo QR payload", "Học viên chuyển khoản", "Quản trị viên kiểm tra và duyệt", "Hệ thống cập nhật gói"],
            ["Chưa thanh toán: trạng thái pending", "Thanh toán sai nội dung: quản trị viên ghi chú"],
        )
        self.usecase_table(
            "Bảng 2.10. Đặc tả use case AI review",
            "AI review",
            "Phân tích kết quả làm bài để chỉ ra điểm mạnh, điểm yếu và hướng ôn tập tiếp theo.",
            "Học viên, Gia sư AI",
            "Học viên nộp quiz hoặc yêu cầu xem nhận xét",
            "Có dữ liệu bài làm hoặc phiên học",
            "Hệ thống hiển thị nhận xét cá nhân hóa",
            [
                "Học viên nộp bài",
                "Backend chấm từng câu",
                "LLM phân tích lỗi theo focus và trình độ",
                "Lưu ai_review_json",
                "Frontend hiển thị điểm mạnh, điểm yếu và gợi ý cải thiện",
            ],
            ["LLM lỗi: hệ thống vẫn hiển thị điểm số và giải thích có sẵn trong câu hỏi"],
        )
        self.usecase_table(
            "Bảng 2.11. Đặc tả use case Thi nâng cấp",
            "Thi nâng cấp",
            "Đánh giá điều kiện chuyển trình độ dựa trên bài kiểm tra tổng hợp theo CEFR.",
            "Học viên, quản trị viên",
            "Học viên chọn bài thi nâng cấp hoặc admin cấp quyền bài thi",
            "Học viên đã có trình độ hiện tại trong hồ sơ",
            "Hệ thống ghi nhận kết quả và đề xuất nâng cấp nếu đạt điều kiện",
            [
                "Hệ thống tạo bài thi theo trình độ kế tiếp",
                "Học viên làm bài",
                "Hệ thống chấm điểm và kiểm tra ngưỡng pass",
                "AI nhận xét kỹ năng cần cải thiện",
                "Trình độ có thể được cập nhật theo chính sách quản trị",
            ],
            ["Không đạt ngưỡng: giữ nguyên trình độ và gợi ý bộ ôn tập"],
        )
        self.usecase_table(
            "Bảng 2.12. Đặc tả use case Quản lý người dùng",
            "Quản lý người dùng",
            "Quản trị viên xem, sửa, xoá mềm hoặc thay đổi thông tin học viên.",
            "Quản trị viên",
            "Quản trị viên mở màn Người dùng",
            "Người dùng có role admin",
            "Thông tin tài khoản được cập nhật trong hệ thống",
            [
                "Admin xem danh sách người dùng",
                "Admin lọc/tìm kiếm tài khoản",
                "Admin sửa tên, role, trình độ hoặc gói học",
                "Hệ thống kiểm tra quyền",
                "Hệ thống lưu thay đổi và trả danh sách mới",
            ],
            ["Không có quyền admin: API trả 403", "Dữ liệu không hợp lệ: hiển thị lỗi validate"],
        )
        self.heading("2.2.3. Phân tích use case tiêu biểu", 3, "h_2_2_3")
        self.image(DIAGRAM_DIR / "vopc_login.png", 6.2, "Hình 2.2. VOPC use case Đăng nhập")
        self.image(DIAGRAM_DIR / "seq_login.png", 6.5, "Hình 2.3. Sơ đồ trình tự use case Đăng nhập")
        self.image(DIAGRAM_DIR / "vopc_practice.png", 6.2, "Hình 2.4. VOPC use case Bắt đầu phiên Gia sư AI")
        self.image(DIAGRAM_DIR / "seq_practice.png", 6.5, "Hình 2.5. Sơ đồ trình tự phiên Gia sư AI")
        self.image(DIAGRAM_DIR / "vopc_chat_quiz.png", 6.2, "Hình 2.6. VOPC use case Làm quiz trong khung chat")
        self.image(DIAGRAM_DIR / "seq_chat_quiz.png", 6.5, "Hình 2.7. Sơ đồ trình tự quiz trong chat")
        self.image(DIAGRAM_DIR / "vopc_admin_quiz.png", 6.2, "Hình 2.8. VOPC use case Quản lý quiz")
        self.image(DIAGRAM_DIR / "seq_admin_quiz.png", 6.5, "Hình 2.9. Sơ đồ trình tự quản lý quiz")
        self.image(DIAGRAM_DIR / "vopc_payment.png", 6.2, "Hình 2.10. VOPC use case Thanh toán QR")
        self.image(DIAGRAM_DIR / "seq_payment.png", 6.5, "Hình 2.11. Sơ đồ trình tự thanh toán QR")

        self.heading("2.3. Thiết kế hệ thống", 2, "h_2_3")
        self.heading("2.3.1. Thiết kế kiến trúc", 3, "h_2_3_1")
        self.para(
            "Hệ thống được chia thành các service độc lập: frontend phục vụ giao diện, API xử lý nghiệp vụ, LiveKit truyền thông thời gian thực, "
            "agent worker xử lý hội thoại, PostgreSQL lưu dữ liệu và mem0 lưu memory cá nhân hóa. Cách tách service này giúp việc phát triển, triển khai "
            "và kiểm thử từng phần rõ ràng hơn."
        )
        self.heading("2.3.2. Thiết kế cơ sở dữ liệu", 3, "h_2_3_2")
        self.image(DIAGRAM_DIR / "erd.png", 6.5, "Hình 2.12. Thiết kế cơ sở dữ liệu chính")
        self.db_tables()
        self.heading("2.3.3. Thiết kế giao diện", 3, "h_2_3_3")
        self.image(DIAGRAM_DIR / "wire_landing.png", 6.0, "Hình 2.13. Thiết kế giao diện Landing page")
        self.image(DIAGRAM_DIR / "wire_practice.png", 6.0, "Hình 2.14. Thiết kế giao diện Gia sư AI")
        self.image(DIAGRAM_DIR / "wire_quiz.png", 6.0, "Hình 2.15. Thiết kế giao diện Kho quiz")
        self.image(DIAGRAM_DIR / "wire_admin.png", 6.0, "Hình 2.16. Thiết kế giao diện Quản trị")
        self.doc.add_page_break()

    def db_tables(self):
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã người dùng", "UUID", "Primary key"],
                ["email", "Email đăng nhập", "String(255)", "Unique, not null"],
                ["password_hash", "Mật khẩu đã băm", "String(255)", "Not null"],
                ["name", "Tên hiển thị", "String(255)", ""],
                ["cefr_level", "Trình độ CEFR hiện tại", "String(2)", "Default B1"],
                ["role", "Vai trò learner/admin", "String(32)", "Default learner"],
                ["subscription_plan", "Gói học", "String(32)", "Default free"],
            ],
            "Bảng 2.13. Mô tả bảng users",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã phiên học", "UUID", "Primary key"],
                ["user_id", "Mã học viên", "UUID", "Foreign key"],
                ["room_name", "Tên phòng LiveKit", "String(255)", ""],
                ["topic", "Chủ đề luyện tập", "String(255)", ""],
                ["level", "Trình độ tại phiên", "String(2)", ""],
                ["total_turns", "Số lượt hội thoại", "Integer", ""],
                ["grammar_score", "Điểm ngữ pháp", "Integer", "Nullable"],
                ["stats_json", "Thống kê chi tiết", "JSON", "Nullable"],
            ],
            "Bảng 2.14. Mô tả bảng practice_sessions",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã lỗi", "UUID", "Primary key"],
                ["session_id", "Phiên học phát sinh lỗi", "UUID", "Foreign key"],
                ["user_id", "Mã học viên", "UUID", "Foreign key"],
                ["error_type", "Loại lỗi", "String(50)", "grammar/vocabulary/..."],
                ["original_text", "Câu ban đầu", "Text", ""],
                ["corrected_text", "Câu đã sửa", "Text", ""],
                ["explanation", "Giải thích", "Text", "Nullable"],
            ],
            "Bảng 2.15. Mô tả bảng error_logs",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã refresh token", "UUID", "Primary key"],
                ["user_id", "Mã người dùng", "UUID", "Foreign key"],
                ["jti", "Định danh token", "String(64)", "Unique"],
                ["expires_at", "Thời điểm hết hạn", "DateTime", "Not null"],
                ["revoked_at", "Thời điểm thu hồi", "DateTime", "Nullable"],
                ["created_at", "Ngày tạo", "DateTime", "Server default"],
            ],
            "Bảng 2.16. Mô tả bảng refresh_tokens",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã bộ đề", "UUID", "Primary key"],
                ["created_by", "Admin tạo bộ đề", "UUID", "Foreign key"],
                ["title", "Tên bộ đề", "String(255)", "Not null"],
                ["source", "Nguồn dữ liệu", "String(32)", "manual/open_source/book"],
                ["source_url", "URL nguồn", "Text", "Nullable"],
                ["license", "Thông tin giấy phép", "Text", "Nullable"],
                ["topic", "Chủ đề", "String(255)", ""],
                ["level", "Trình độ", "String(2)", ""],
            ],
            "Bảng 2.17. Mô tả bảng quiz_sets",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã quiz", "UUID", "Primary key"],
                ["user_id", "Người tạo", "UUID", "Foreign key"],
                ["quiz_set_id", "Bộ đề", "UUID", "Foreign key, nullable"],
                ["title", "Tiêu đề quiz", "String(255)", "Not null"],
                ["topic", "Chủ đề", "String(255)", ""],
                ["level", "Trình độ", "String(2)", ""],
                ["source", "Nguồn tạo", "String(32)", "ai/manual/imported/..."],
                ["questions_json", "Danh sách câu hỏi", "JSON", "Not null"],
            ],
            "Bảng 2.18. Mô tả bảng quizzes",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã lượt làm bài", "UUID", "Primary key"],
                ["quiz_id", "Mã quiz", "UUID", "Foreign key"],
                ["user_id", "Mã học viên", "UUID", "Foreign key"],
                ["answers_json", "Đáp án của học viên", "JSON", "Not null"],
                ["result_json", "Kết quả từng câu", "JSON", "Not null"],
                ["ai_review_json", "Nhận xét AI", "JSON", "Nullable"],
                ["score", "Điểm số", "Integer", ""],
                ["correct_count", "Số câu đúng", "Integer", ""],
            ],
            "Bảng 2.19. Mô tả bảng quiz_attempts",
            [3.2, 5.0, 3.2, 3.4],
        )
        self.table(
            [
                ["Tên thuộc tính", "Diễn giải", "Kiểu dữ liệu", "Ràng buộc"],
                ["id", "Mã yêu cầu", "UUID", "Primary key"],
                ["user_id", "Mã học viên", "UUID", "Foreign key"],
                ["plan", "Gói đăng ký", "String(32)", "Not null"],
                ["amount_vnd", "Số tiền VNĐ", "Integer", "Not null"],
                ["status", "Trạng thái", "String(32)", "pending/approved/rejected"],
                ["qr_payload", "Nội dung QR", "Text", "Not null"],
                ["admin_note", "Ghi chú quản trị", "Text", "Nullable"],
            ],
            "Bảng 2.20. Mô tả bảng payment_requests",
            [3.2, 5.0, 3.2, 3.4],
        )

    def chapter3(self):
        self.heading("CHƯƠNG 3. CÀI ĐẶT VÀ KẾT QUẢ", 1, "h_ch3")
        self.heading("3.1. Cài đặt hệ thống", 2, "h_3_1")
        self.heading("3.1.1. Yêu cầu cấu hình", 3, "h_3_1_1")
        self.table(
            [
                ["Thành phần", "Công nghệ", "Vai trò"],
                ["Frontend", "React, TypeScript, Vite, Tailwind", "Giao diện học viên và quản trị viên"],
                ["Backend API", "FastAPI, SQLAlchemy, Pydantic", "Xác thực, quiz, billing, token"],
                ["Realtime", "LiveKit, WebRTC, TURN", "Truyền âm thanh, data channel, widget"],
                ["AI Agent", "LiveKit Agents, LangGraph, STT/TTS", "Hội thoại, sửa lỗi, tạo bài tập"],
                ["Database", "PostgreSQL, pgvector", "Lưu dữ liệu nghiệp vụ và memory"],
                ["Deploy", "Docker Compose, Vercel", "Triển khai backend/service và frontend"],
            ],
            "Bảng 3.1. Thành phần triển khai hệ thống",
            [3.5, 5.0, 6.0],
        )
        self.heading("3.1.2. Hướng dẫn triển khai", 3, "h_3_1_2")
        self.table(
            [
                ["Nhóm cấu hình", "Biến/cổng", "Mục đích"],
                ["API", "DATABASE_URL, JWT_SECRET, REFRESH_SECRET", "Kết nối DB và ký token"],
                ["LiveKit", "LIVEKIT_URL, LIVEKIT_API_KEY, LIVEKIT_API_SECRET", "Tạo room và cấp token WebRTC"],
                ["Frontend", "VITE_API_BASE_URL, VITE_LIVEKIT_URL", "Kết nối giao diện tới backend và LiveKit"],
                ["AI", "GEMINI_API_KEY/GROQ_API_KEY", "Gọi mô hình ngôn ngữ lớn"],
                ["Network", "5173, 8080, 7880, 7881, 3478", "Frontend, API, LiveKit, TURN"],
            ],
            "Bảng 3.2. Cấu hình môi trường",
            [3.5, 5.2, 5.8],
        )
        self.table(
            [
                ["Nhóm API", "Endpoint tiêu biểu", "Chức năng"],
                ["Auth", "/auth/register, /auth/login, /auth/refresh", "Tạo tài khoản, đăng nhập, xoay vòng token"],
                ["Password reset", "/auth/password/forgot, /auth/password/reset", "Gửi mã xác thực email và đặt lại mật khẩu"],
                ["LiveKit token", "/token", "Cấp token vào phòng Gia sư AI và kiểm tra quota chat"],
                ["Quizzes", "/quizzes, /quizzes/{id}/submit", "Danh sách quiz, làm bài, lưu kết quả"],
                ["Admin quizzes", "/quizzes/admin/*", "Tạo, sửa, xoá, import nguồn quiz"],
                ["Billing", "/billing/payment-requests", "Tạo và duyệt yêu cầu thanh toán QR"],
                ["Admin users", "/admin/users", "Quản lý role, trình độ và gói học của người dùng"],
            ],
            "Bảng 3.3. Các API chính của hệ thống",
            [3.2, 5.2, 6.0],
        )
        self.table(
            [
                ["Bước", "Thành phần xử lý", "Kết quả"],
                ["1", "Frontend + LiveKit", "Học viên gửi audio hoặc text vào phòng"],
                ["2", "STT/VAD", "Chuyển giọng nói thành transcript, gom đoạn để giảm duplicate"],
                ["3", "LangGraph assess node", "Phân tích lỗi, đánh giá chất lượng và trình độ gợi ý"],
                ["4", "Router node", "Quyết định phản hồi thường, sửa lỗi, đổi chủ đề hoặc mở bài tập"],
                ["5", "Respond/Correct node", "Sinh câu trả lời phù hợp CEFR và chủ đề"],
                ["6", "Widget builder", "Tạo quiz/recap/mistake notebook khi đúng intent"],
                ["7", "TTS + data channel", "Trả giọng nói, transcript và widget về frontend"],
            ],
            "Bảng 3.4. Luồng xử lý AI trong phiên Gia sư",
            [1.4, 4.8, 8.0],
        )
        self.para(
            "Khi chạy local, frontend hoạt động ở cổng 5173, API ở cổng 8080 và LiveKit ở cổng 7880. Khi deploy, frontend có thể chạy trên Vercel "
            "và gọi backend công khai/được reverse proxy; các cổng nội bộ có thể được truy cập qua SSH port-forward trong giai đoạn kiểm thử."
        )
        self.heading("3.2. Kết quả thực hiện", 2, "h_3_2")
        self.heading("3.2.1. Giao diện người dùng", 3, "h_3_2_1")
        self.image(UI_DIR / "landing.png", 6.4, "Hình 3.1. Giao diện Landing page")
        self.image(UI_DIR / "login.png", 6.0, "Hình 3.2. Giao diện đăng nhập")
        self.image(UI_DIR / "learner_home.png", 6.4, "Hình 3.3. Giao diện trang học tập")
        self.image(UI_DIR / "practice_ready.png", 6.2, "Hình 3.4. Giao diện Gia sư AI")
        self.image(UI_DIR / "learner_quizzes.png", 6.4, "Hình 3.5. Giao diện kho quiz")
        self.image(UI_DIR / "billing.png", 6.4, "Hình 3.6. Giao diện thanh toán QR")
        self.heading("3.2.2. Giao diện quản trị", 3, "h_3_2_2")
        self.para(
            "Luồng quản trị được tách khỏi luồng học viên bằng role admin. Quản trị viên truy cập các màn hình tổng quan, quản lý người dùng, quản lý quiz "
            "và quản lý thanh toán. Học viên không thể truy cập các route /admin do frontend và backend đều kiểm tra vai trò."
        )
        self.heading("3.2.3. Kết quả xử lý nghiệp vụ AI", 3, "h_3_2_3")
        self.para(
            "Agent sử dụng prompt theo trình độ CEFR, phân tích lỗi của học viên, quyết định khi nào phản hồi thường, khi nào sửa lỗi rõ ràng và khi nào "
            "gửi widget. Khi học viên yêu cầu làm bài tập, agent ưu tiên tạo quiz trắc nghiệm trong chat; nếu học viên nói rõ luyện nghe hoặc luyện nói thì "
            "hệ thống mở widget tương ứng."
        )
        self.bullets([
            "STT chuyển lời nói của học viên thành transcript, transcript được gom để tránh duplicate tin nhắn.",
            "TTS phát giọng nói của agent giúp phiên học giống hội thoại trực tuyến.",
            "Widget quiz trong chat hiển thị câu hỏi, đáp án, chấm điểm và review ngay trong phiên học.",
            "Kết quả quiz ngoài kho được lưu vào quiz_attempts để phục vụ dashboard và AI review.",
            "Giới hạn gói học kiểm soát số lượt chat và quiz theo ngày, hỗ trợ upsell sang Plus/Ultra.",
        ])
        self.heading("3.3. Kiểm thử", 2, "h_3_3")
        self.heading("3.3.1. Kế hoạch kiểm thử", 3, "h_3_3_1")
        self.para(
            "Kiểm thử tập trung vào các luồng quan trọng có ảnh hưởng trực tiếp tới trải nghiệm học viên: xác thực, phiên Gia sư AI, quiz, AI review, "
            "phân quyền admin và thanh toán QR. Ngoài kiểm thử chức năng, hệ thống được kiểm tra thêm khả năng kết nối WebRTC và giới hạn quota."
        )
        self.heading("3.3.2. Lịch trình kiểm thử", 3, "h_3_3_2")
        self.table(
            [
                ["Mốc công việc", "Sản phẩm", "Thời gian", "Kết quả"],
                ["Lập test plan", "Danh sách kịch bản", "1 ngày", "Hoàn thành"],
                ["Kiểm thử auth/admin", "Test case", "1 ngày", "Hoàn thành"],
                ["Kiểm thử Gia sư AI", "Test case", "2 ngày", "Hoàn thành"],
                ["Kiểm thử quiz/widget", "Test case", "2 ngày", "Hoàn thành"],
                ["Kiểm thử billing/quota", "Test case", "1 ngày", "Hoàn thành"],
                ["Ghi nhận và sửa lỗi", "Bug fix", "Liên tục", "Hoàn thành ở mức MVP"],
            ],
            "Bảng 3.5. Lịch trình kiểm thử",
            [4.2, 3.8, 3.0, 4.0],
        )
        self.heading("3.3.3. Kết quả kiểm thử", 3, "h_3_3_3")
        self.table(
            [
                ["Mã", "Chức năng", "Kịch bản", "Kết quả mong muốn", "Trạng thái"],
                ["TC01", "Đăng ký/đăng nhập", "Tạo tài khoản, đăng nhập, refresh phiên", "Người dùng vào đúng màn hình", "Đạt"],
                ["TC02", "Phân quyền", "Learner truy cập /admin", "Bị chặn và chuyển hướng", "Đạt"],
                ["TC03", "Gia sư AI", "Bắt đầu phiên và gửi tin nhắn", "Agent trả lời, không duplicate", "Đạt"],
                ["TC04", "Voice chat", "Nói qua micro", "Có transcript và agent phản hồi bằng giọng nói", "Đạt"],
                ["TC05", "Quiz trong chat", "Yêu cầu làm bài tập", "Hiển thị widget trắc nghiệm", "Đạt"],
                ["TC06", "Kho quiz", "Mở bộ đề, làm bài, nộp bài", "Tính điểm và lưu attempt", "Đạt"],
                ["TC07", "AI review", "Xem kết quả sau quiz", "Có điểm mạnh, điểm yếu, gợi ý cải thiện", "Đạt"],
                ["TC08", "Thanh toán QR", "Tạo yêu cầu mua gói", "Sinh QR và pending request", "Đạt"],
                ["TC09", "Hết lượt", "Dùng quá quota gói free", "Chặn chat/quiz và gợi ý mua gói", "Đạt một phần"],
            ],
            "Bảng 3.6. Kết quả kiểm thử chức năng",
            [1.4, 3.0, 4.0, 4.2, 2.0],
        )
        self.heading("3.3.4. Điều kiện chấp nhận kiểm thử", 3, "h_3_3_4")
        self.bullets([
            "Các luồng chính của học viên chạy được từ đăng nhập, luyện nói, làm quiz đến xem kết quả.",
            "Admin có thể quản lý quiz, người dùng và thanh toán mà học viên không truy cập được.",
            "Phiên LiveKit kết nối ổn định trong môi trường local/forward port hoặc môi trường deploy.",
            "Các lỗi từng gặp như duplicate tin nhắn, agent chen ngang quá sớm và widget quiz không xuất hiện đã được xử lý ở mức MVP.",
            "Các rủi ro còn lại gồm chất lượng bộ dữ liệu quiz, độ chính xác đánh giá trình độ và tích hợp thanh toán tự động với ngân hàng.",
        ])
        self.doc.add_page_break()

    def conclusion(self):
        self.heading("KẾT LUẬN", 1, "h_conclusion")
        self.para(
            "Đồ án đã xây dựng được hệ thống Gia sư AI tiếng Anh có đầy đủ các chức năng cốt lõi của một sản phẩm học tập MVP: xác thực, phân quyền, "
            "luyện nói với agent thời gian thực, TTS/STT, sửa lỗi, quiz theo bộ, quiz trong chat, AI review, thi nâng cấp, quản trị hệ thống và thanh toán QR."
        )
        self.para(
            "So với chatbot thông thường, hệ thống hướng tới trải nghiệm học tập có cấu trúc hơn: học viên được gắn với trình độ CEFR, nội dung quiz bị giới hạn "
            "theo trình độ, lỗi sai được ghi nhận và agent có thể chuyển từ hội thoại sang bài tập ngay trong khung chat."
        )
        self.para(
            "Hướng phát triển tiếp theo là mở rộng dữ liệu quiz từ nguồn có bản quyền rõ ràng, chuẩn hóa bài thi nâng cấp theo kỹ năng nghe-nói-đọc-viết, tích hợp "
            "thanh toán tự động qua cổng thanh toán như payOS và tăng độ chính xác của cá nhân hóa dựa trên lịch sử học tập dài hạn."
        )
        self.doc.add_page_break()
        self.heading("TÀI LIỆU THAM KHẢO", 1, "h_refs")
        refs = [
            "FastAPI Documentation. https://fastapi.tiangolo.com/",
            "React Documentation. https://react.dev/",
            "LiveKit Documentation. https://docs.livekit.io/",
            "LangGraph Documentation. https://langchain-ai.github.io/langgraph/",
            "Council of Europe, Common European Framework of Reference for Languages (CEFR).",
            "PostgreSQL Documentation. https://www.postgresql.org/docs/",
            "Vite Documentation. https://vitejs.dev/",
        ]
        for idx, ref in enumerate(refs, 1):
            self.para(f"[{idx}] {ref}", first_line=False)

    def add_footer_page_numbers(self):
        for section in self.doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_field(p, "PAGE", "")

    def build(self):
        self.cover()
        self.front_matter()
        self.intro()
        self.chapter1()
        self.chapter2()
        self.chapter3()
        self.conclusion()
        self.add_footer_page_numbers()
        self.doc.save(OUT_DOCX)
        set_update_fields(OUT_DOCX)


def main():
    generate_assets()
    builder = Builder()
    builder.build()
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
