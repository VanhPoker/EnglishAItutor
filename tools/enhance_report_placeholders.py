from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/nguyenvietanh/Main-GKE/EnglishAItutor")
SOURCE = ROOT / "EnglishAItutor_DATN.docx"
OUTPUT = ROOT / "EnglishAItutor_DATN_v2_placeholders.docx"
ASSET_DIR = ROOT / "report_assets_expanded" / "placeholders"


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_centered(draw, xy, text, fnt, fill):
    x1, y1, x2, y2 = xy
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=8, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(
        (x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2),
        text,
        font=fnt,
        fill=fill,
        spacing=8,
        align="center",
    )


def dashed_rect(draw, rect, color, width=4, dash=22, gap=14):
    x1, y1, x2, y2 = rect
    for x in range(x1, x2, dash + gap):
        draw.line([(x, y1), (min(x + dash, x2), y1)], fill=color, width=width)
        draw.line([(x, y2), (min(x + dash, x2), y2)], fill=color, width=width)
    for y in range(y1, y2, dash + gap):
        draw.line([(x1, y), (x1, min(y + dash, y2))], fill=color, width=width)
        draw.line([(x2, y), (x2, min(y + dash, y2))], fill=color, width=width)


def make_placeholder(slug: str, title: str, detail: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / f"{slug}.png"
    img = Image.new("RGB", (1600, 900), "#f7fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1599, 899], fill="#f7fafc")
    dashed_rect(draw, (60, 60, 1540, 840), "#64748b", width=5)
    draw.rectangle([90, 90, 1510, 810], outline="#cbd5e1", width=2)
    draw_centered(
        draw,
        (160, 210, 1440, 530),
        f"VÙNG CHÈN ẢNH GIAO DIỆN THẬT\n{title}",
        font(48, True),
        "#1e293b",
    )
    draw_centered(
        draw,
        (210, 560, 1390, 700),
        detail,
        font(30),
        "#475569",
    )
    draw_centered(
        draw,
        (260, 715, 1340, 780),
        "Sau khi chụp màn hình, thay ảnh placeholder này bằng ảnh giao diện thực tế.",
        font(24),
        "#64748b",
    )
    img.save(path)
    return path


def paragraph_after(paragraph, text="", style=None):
    new_para = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new_para._p)
    return new_para


def paragraph_before(paragraph, text="", style=None):
    new_para = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addprevious(new_para._p)
    return new_para


def format_body(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def format_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.italic = True


def add_picture_block_before(target, image_path: Path, caption: str):
    img_para = paragraph_before(target, "")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(image_path), width=Cm(15.6))
    cap = paragraph_before(target, caption, style="CaptionVN")
    format_caption(cap)
    spacer = paragraph_before(target, "")
    spacer.paragraph_format.space_after = Pt(8)


def add_page_break_before(target):
    p = paragraph_before(target, "")
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def find_paragraph(doc: Document, exact: str):
    for para in doc.paragraphs:
        if para.text.strip() == exact:
            return para
    raise ValueError(f"Cannot find paragraph: {exact}")


def remove_between(doc: Document, start_heading: str, end_heading: str):
    start = find_paragraph(doc, start_heading)
    end = find_paragraph(doc, end_heading)
    current = start._p.getnext()
    while current is not None and current is not end._p:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt
    return start, end


def insert_after_chain(anchor, items):
    current = anchor
    for text, style in items:
        p = paragraph_after(current, text, style=style)
        if style in (None, "Normal"):
            format_body(p)
        current = p
    return current


def rebuild_conclusion_and_references(doc: Document):
    conclusion, references = remove_between(doc, "KẾT LUẬN", "TÀI LIỆU THAM KHẢO")
    insert_after_chain(
        conclusion,
        [
            ("1. Kết quả đạt được", "Heading 2"),
            (
                "Đề tài đã xây dựng được phần mềm EnglishAItutor theo mô hình full-stack, đáp ứng mục tiêu hỗ trợ người học luyện giao tiếp tiếng Anh thông qua hội thoại với AI. Hệ thống cho phép người dùng đăng ký, đăng nhập, chọn trình độ CEFR, chọn chủ đề luyện tập và tham gia phiên học bằng giọng nói hoặc văn bản. Các thành phần frontend, backend API và agent worker được tách riêng, giúp cấu trúc mã nguồn rõ ràng và thuận lợi cho mở rộng.",
                "Normal",
            ),
            (
                "Ở phía backend, hệ thống sử dụng FastAPI để cung cấp API xác thực, quản lý phiên học, cấp token LiveKit và thống kê dashboard. Phần agent hội thoại được tổ chức bằng LangGraph, trong đó các bước đánh giá đầu vào, chuẩn bị ngữ cảnh, phản hồi, sửa lỗi và lưu bộ nhớ được mô hình hóa thành các node riêng. Cách thiết kế này giúp luồng xử lý có tính kiểm soát, dễ theo dõi và dễ bổ sung nghiệp vụ mới.",
                "Normal",
            ),
            (
                "Ở phía frontend, EnglishAItutor cung cấp các màn hình chính gồm đăng nhập/đăng ký, trang thiết lập phiên học, phòng luyện tập, dashboard tiến độ và trang quản trị người dùng. Giao diện tập trung vào thao tác học tập, hiển thị trạng thái kết nối, transcript hội thoại, phản hồi sửa lỗi và các chỉ số theo dõi quá trình luyện tập. Báo cáo cũng đã bổ sung mockup và vùng chèn ảnh giao diện thực tế để hoàn thiện minh chứng sản phẩm sau khi chạy demo.",
                "Normal",
            ),
            ("2. Hạn chế của đề tài", "Heading 2"),
            (
                "Do phạm vi thời gian của đồ án, hệ thống hiện mới tập trung vào luồng luyện tập cốt lõi và chưa triển khai đầy đủ các bài học có cấu trúc theo từng kỹ năng. Chất lượng phản hồi của AI vẫn phụ thuộc vào mô hình ngôn ngữ, dịch vụ STT/TTS và chất lượng âm thanh đầu vào. Một số tiêu chí như phát âm theo âm vị, ngữ điệu, tốc độ nói và mức độ tự nhiên của hội thoại mới được đánh giá ở mức chức năng, chưa có bộ đo chuyên sâu.",
                "Normal",
            ),
            (
                "Hệ thống cũng cần bổ sung thêm kiểm thử tự động cho các luồng real-time sử dụng LiveKit, kiểm thử tải khi nhiều người học tham gia đồng thời và cơ chế giám sát vận hành trong môi trường production. Phần dashboard đã có các chỉ số cơ bản nhưng chưa có phân tích xu hướng học tập dài hạn hoặc đề xuất lộ trình cá nhân hóa theo mục tiêu cụ thể của từng người học.",
                "Normal",
            ),
            ("3. Hướng phát triển", "Heading 2"),
            (
                "Trong giai đoạn tiếp theo, EnglishAItutor có thể được mở rộng theo hướng xây dựng lộ trình học cá nhân hóa dựa trên CEFR, mục tiêu học tập và lịch sử lỗi của người dùng. Hệ thống có thể bổ sung ngân hàng bài học, chủ đề hội thoại theo tình huống thực tế, bài luyện phát âm, bài luyện phản xạ phỏng vấn và cơ chế đánh giá tiến bộ theo tuần hoặc theo tháng.",
                "Normal",
            ),
            (
                "Về kỹ thuật, hệ thống nên bổ sung pipeline kiểm thử end-to-end cho frontend, API và agent; triển khai quan sát hệ thống bằng logging, tracing và dashboard vận hành; tối ưu chi phí gọi LLM bằng cache, routing model và kiểm soát prompt. Khi đưa vào sử dụng thực tế, cần hoàn thiện HTTPS, backup dữ liệu, phân quyền quản trị, chính sách bảo mật thông tin học tập và bộ tiêu chí đánh giá chất lượng phản hồi của AI.",
                "Normal",
            ),
        ],
    )
    add_page_break_before(references)

    _, appendix = remove_between(doc, "TÀI LIỆU THAM KHẢO", "PHỤ LỤC")
    refs = [
        "FastAPI. FastAPI Documentation. https://fastapi.tiangolo.com/ (truy cập ngày 24/04/2026).",
        "LiveKit. LiveKit Documentation. https://docs.livekit.io/ (truy cập ngày 24/04/2026).",
        "LangChain. LangGraph Documentation. https://langchain-ai.github.io/langgraph/ (truy cập ngày 24/04/2026).",
        "LangChain. LangChain Documentation. https://python.langchain.com/ (truy cập ngày 24/04/2026).",
        "React. React Documentation. https://react.dev/ (truy cập ngày 24/04/2026).",
        "Vite. Vite Documentation. https://vite.dev/ (truy cập ngày 24/04/2026).",
        "Tailwind Labs. Tailwind CSS Documentation. https://tailwindcss.com/docs (truy cập ngày 24/04/2026).",
        "PostgreSQL Global Development Group. PostgreSQL Documentation. https://www.postgresql.org/docs/ (truy cập ngày 24/04/2026).",
        "pgvector. Open-source vector similarity search for PostgreSQL. https://github.com/pgvector/pgvector (truy cập ngày 24/04/2026).",
        "Docker. Docker Compose Documentation. https://docs.docker.com/compose/ (truy cập ngày 24/04/2026).",
        "SQLAlchemy. SQLAlchemy Documentation. https://docs.sqlalchemy.org/ (truy cập ngày 24/04/2026).",
        "Alembic. Alembic Documentation. https://alembic.sqlalchemy.org/ (truy cập ngày 24/04/2026).",
        "Pydantic. Pydantic Documentation. https://docs.pydantic.dev/ (truy cập ngày 24/04/2026).",
        "Council of Europe. Common European Framework of Reference for Languages: Learning, teaching, assessment.",
        "OpenAI. Prompt engineering and structured AI application design references. https://platform.openai.com/docs/ (truy cập ngày 24/04/2026).",
        "Google AI. Gemini API Documentation. https://ai.google.dev/gemini-api/docs (truy cập ngày 24/04/2026).",
        "Groq. GroqCloud Documentation. https://console.groq.com/docs (truy cập ngày 24/04/2026).",
        "mem0. Memory layer for AI applications. https://docs.mem0.ai/ (truy cập ngày 24/04/2026).",
        "EnglishAItutor source code. Thư mục mã nguồn /Users/nguyenvietanh/Main-GKE/EnglishAItutor.",
    ]
    current = references
    for index, item in enumerate(refs, start=1):
        p = paragraph_after(current, f"[{index}] {item}", style="Normal")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
        current = p
    add_page_break_before(appendix)


def update_static_toc(doc: Document):
    toc_entries = {
        "LỜI CẢM ƠN": 2,
        "TÓM TẮT ĐỒ ÁN": 3,
        "DANH MỤC HÌNH": 5,
        "DANH MỤC BẢNG": 6,
        "MỞ ĐẦU": 7,
        "CHƯƠNG 1. CƠ SỞ LÝ THUYẾT": 9,
        "CHƯƠNG 2. KHẢO SÁT, PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG": 12,
        "CHƯƠNG 3. THIẾT KẾ GIAO DIỆN VÀ MOCKUP": 31,
        "CHƯƠNG 4. XÂY DỰNG HỆ THỐNG": 41,
        "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ": 51,
        "KẾT LUẬN": 58,
        "TÀI LIỆU THAM KHẢO": 60,
        "PHỤ LỤC": 61,
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        for title, page in toc_entries.items():
            if text.startswith(title + " "):
                page_text = str(page)
                dot_count = max(5, 78 - len(title) - len(page_text))
                para.text = f"{title} {'.' * dot_count} {page_text}"
                para.paragraph_format.left_indent = Cm(0)
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                break


def add_ui_placeholders(doc: Document):
    ch4 = find_paragraph(doc, "CHƯƠNG 4. XÂY DỰNG HỆ THỐNG")
    intro = paragraph_before(ch4, "3.7. Vùng chèn ảnh giao diện thực tế", style="Heading 2")
    p = paragraph_before(
        ch4,
        "Các khung dưới đây được bố trí sẵn theo tỉ lệ ảnh màn hình phổ biến để thay bằng ảnh chụp giao diện thật sau khi chạy demo sản phẩm. Khi hoàn thiện báo cáo, chỉ cần thay placeholder bằng ảnh chụp tương ứng và giữ nguyên caption bên dưới.",
        style="Normal",
    )
    format_body(p)

    placeholders = [
        ("real_login", "LOGIN / REGISTER", "Ảnh đề xuất: màn hình đăng nhập hoặc đăng ký với dữ liệu demo."),
        ("real_home", "HOME / SESSION SETUP", "Ảnh đề xuất: màn chọn CEFR, topic và nút bắt đầu phiên luyện tập."),
        ("real_practice", "PRACTICE ROOM", "Ảnh đề xuất: phiên hội thoại đang kết nối, có transcript và phản hồi tutor."),
        ("real_dashboard", "DASHBOARD", "Ảnh đề xuất: thống kê số phiên, phút luyện, lỗi đã sửa và biểu đồ tiến độ."),
        ("real_admin", "ADMIN USERS", "Ảnh đề xuất: bảng quản trị người dùng, role, CEFR và trạng thái tài khoản."),
        ("real_mobile", "RESPONSIVE / MOBILE", "Ảnh đề xuất: giao diện mobile hoặc tablet nếu cần minh chứng responsive."),
    ]
    for idx, (slug, title, detail) in enumerate(placeholders, start=6):
        path = make_placeholder(slug, title, detail)
        add_picture_block_before(ch4, path, f"Hình 3.{idx}. Vùng chèn ảnh giao diện thực tế - {title.title()}")
    add_page_break_before(ch4)

    fig_list = find_paragraph(doc, "Hình 3.5. Mockup Admin Users")
    current = fig_list
    for idx, (_, title, _) in enumerate(placeholders, start=6):
        p = paragraph_after(current, f"Hình 3.{idx}. Vùng chèn ảnh giao diện thực tế - {title.title()}", style="Normal")
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
        current = p


def main():
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    rebuild_conclusion_and_references(doc)
    add_ui_placeholders(doc)
    update_static_toc(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
